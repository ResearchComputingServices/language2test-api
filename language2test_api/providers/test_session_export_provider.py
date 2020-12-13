from sqlalchemy.sql import text
from flask import json, jsonify, Response, blueprints, request
from language2test_api.extensions import db, ma
from collections import namedtuple
from language2test_api.providers.raw_sql_provider import RawSqlProvider
from language2test_api.providers.test_session_export_helper import TestSessionExportHelper as testhelper
from language2test_api.models.test_session import TestSession, TestSessionSchema
from typing import List
import pandas as pd
from io import BytesIO
from docx import Document
import zipfile

test_schema = TestSessionSchema(many=False)


class TestSessionExportProvider(RawSqlProvider):
    def write_results_into_file(self, sessions, name):
        xlsx = BytesIO()
        with pd.ExcelWriter(xlsx, engine='xlsxwriter') as writer:
            for session_id in [session.id for session in sessions]:
                t_s = TestSession.query.filter_by(id=session_id).first()
                results = test_schema.dump(t_s)
                test_time = testhelper.calculate_time_consumption(results['end_datetime'], results['start_datetime'])

                basic_info = {
                    "User ID": results['user_id'],
                    "Test Info": results['name'],
                    "User Name": results['user']['name'],
                    "First Name": results['user']['first_name'],
                    "Last Name": results['user']['last_name'],
                    "Start Time": results['start_datetime'],
                    "End Time": results['end_datetime'],
                    "Time Consumption": test_time
                }

                if ('results_vocabulary' in results) and ((len(results['results_vocabulary'])) > 0):
                    voc_score = testhelper.calculate_correctly_answered_questions(
                        results['results_vocabulary'][0]['answers'])
                    voc_total = len(results['results_vocabulary'][0]['answers'])
                    voc_grade = testhelper.grade_calculator(voc_score, voc_total)
                    if voc_total > 0:
                        basic_info["Vocabulary Grade"] = voc_grade

                if 'results_rc' in results and ((len(results['results_rc'])) > 0):
                    rc_score = testhelper.calculate_correctly_answered_questions(results['results_rc'][0]['answers'])
                    rc_total = len(results['results_rc'][0]['answers'])
                    rc_grade = testhelper.grade_calculator(rc_score, rc_total)
                    if rc_total > 0:
                        basic_info["Reading Comprehension Grade"] = rc_grade

                if 'results_cloze' in results and ((len(results['results_cloze'])) > 0):
                    cloze_score = testhelper.calculate_correctly_answered_questions(
                        results['results_cloze'][0]['answers'])
                    cloze_total = len(results['results_cloze'][0]['answers'])
                    cloze_grade = testhelper.grade_calculator(cloze_score, cloze_total)
                    if cloze_total > 0:
                        basic_info["Cloze Grade"] = cloze_grade

                if 'results_writing' in results and ((len(results['results_writing'])) > 0):
                    writing_score = testhelper.calculate_correctly_answered_questions_writings(
                        results['results_writing'])
                    writing_total = len(results['results_writing'])
                    writing_grade = testhelper.grade_calculator(writing_score, writing_total)
                    if writing_total > 0:
                        basic_info["Writing Grade"] = writing_grade
                else:
                    basic_info["Writing Grade"] = None

                test_results = results['test']

                if ('test_vocabulary' in test_results) and (len(test_results['test_vocabulary']) > 0):
                    voc_results = test_results["test_vocabulary"]
                    voc_infos = []
                    for i in range(len(results['results_vocabulary'][0]['answers'])):
                        answer = results['results_vocabulary'][0]['answers'][i]['text']
                        if answer is None:
                            test_result = "Empty Answer"
                        elif answer == voc_results[i]['options'][int(voc_results[i]['correct']) - 1]['text']:
                            test_result = "Correct"
                        else:
                            test_result = "Incorrect"
                        test_voc_time = testhelper.calculate_time_consumption(
                            results['results_vocabulary'][0]['answers'][i]['end_time'],
                            results['results_vocabulary'][0]['answers'][i]['start_time'])
                        voc_infos.append({
                            "Vocabulary": str(i + 1),
                            "Test Type": voc_results[i]['type'],
                            "Word": voc_results[i]['word'],
                            "Option 1": voc_results[i]['options'][0]['text'],
                            "Option 2": voc_results[i]['options'][1]['text'],
                            "Option 3": voc_results[i]['options'][2]['text'],
                            "Option 4": voc_results[i]['options'][3]['text'],
                            "Student Answer": answer,
                            "Correct Answer": voc_results[i]['options'][int(voc_results[i]['correct']) - 1]['text'],
                            "Difficulty": voc_results[i]['difficulty'],
                            "Test Result": test_result,
                            "Time Consumption": test_voc_time
                        })

                if 'test_rc' in test_results and (len(test_results['test_rc']) > 0):
                    rc_results = test_results["test_rc"]
                    rc_infos = []
                    for i in range(len(rc_results)):
                        rc_questions = test_results["test_rc"][i]['questions']
                        for j in range(len(rc_questions)):
                            answer = results['results_rc'][i]['answers'][j]['text']
                            if answer is None:
                                test_result = "Empty Answer"
                            elif answer == rc_questions[j]['options'][int(rc_questions[j]['correct']) - 1]['text']:
                                test_result = "Correct"
                            else:
                                test_result = "Incorrect"
                            test_rc_time = testhelper.calculate_time_consumption(
                                results['results_rc'][i]['answers'][j]['end_time'],
                                results['results_rc'][i]['answers'][j]['start_time'])
                            rc_infos.append({
                                "RC": str(j + 1),
                                "Question": rc_questions[j]['text'],
                                "Option 1": rc_questions[j]['options'][0]['text'],
                                "Option 2": rc_questions[j]['options'][1]['text'],
                                "Option 3": rc_questions[j]['options'][2]['text'],
                                "Option 4": rc_questions[j]['options'][3]['text'],
                                "User Answer": results['results_rc'][i]['answers'][j]['text'],
                                "Correct Answer": rc_questions[j]['options'][int(rc_questions[j]['correct']) - 1][
                                    'text'],
                                "Test Result": test_result,
                                "Time Consumption": test_rc_time
                            })

                if 'test_cloze' in test_results and (len(test_results['test_cloze']) > 0):
                    cloze_results = test_results["test_cloze"]

                    cloze_infos = []
                    for i in range(len(cloze_results)):
                        cloze_questions = test_results["test_cloze"][i]['questions']
                        for j in range(len(cloze_questions)):
                            answer = results['results_cloze'][i]['answers'][j]['text']
                            if answer is None:
                                test_result = "Empty Answer"
                            elif answer == cloze_questions[j]['options'][int(cloze_questions[j]['correct']) - 1][
                                'text']:
                                test_result = "Correct"
                            else:
                                test_result = "Incorrect"
                            test_cloze_time = testhelper.calculate_time_consumption(
                                results['results_cloze'][i]['answers'][j]['end_time'],
                                results['results_cloze'][i]['answers'][j]['start_time'])
                            cloze_infos.append({
                                "Cloze": str(j + 1),
                                "Question": cloze_questions[j]['text'],
                                "Option 1": cloze_questions[j]['options'][0]['text'],
                                "Option 2": cloze_questions[j]['options'][1]['text'],
                                "Option 3": cloze_questions[j]['options'][2]['text'],
                                "Option 4": cloze_questions[j]['options'][3]['text'],
                                "User Answer": results['results_cloze'][i]['answers'][j]['text'],
                                "Correct Answer": cloze_questions[j]['options'][int(cloze_questions[j]['correct']) - 1][
                                    'text'],
                                "Test Result": test_result,
                                "Time Consumption": test_cloze_time
                            })

                if 'test_writing' in test_results and (len(test_results['test_writing']) > 0):
                    writing_results = test_results["test_writing"]
                    writing_infos = []
                    for i in range(len(writing_results)):
                        test_writing_time = testhelper.calculate_time_consumption(
                            results['results_writing'][i]['answer']['end_time'],
                            results['results_writing'][i]['answer']['start_time'])
                        writing_infos.append({
                            "Writing": str(i + 1),
                            "Test name": writing_results[i]['name'],
                            "Question": writing_results[i]['question'],
                            "Word Limit": writing_results[i]['word_limit'],
                            "Time Limit": writing_results[i]['time_limit'],
                            "Essay": results['results_writing'][i]['answer']['text'],
                            "Time Consumption": test_writing_time
                        })

                if ('results_vocabulary' in results) and ((len(results['results_vocabulary'])) > 0):
                    if voc_total > 0:
                        pd.DataFrame(voc_infos).to_excel(writer, sheet_name="Vocabulary Test Info #" + str(session_id),
                                                         index=False)
                        pd.DataFrame([voc_grade], index=["Grade"]).to_excel(writer, startrow=1 + len(voc_infos),
                                                                            sheet_name="Vocabulary Test Info #" + str(
                                                                                session_id),
                                                                            header=False)
                        workbook = writer.book
                        format = workbook.add_format()
                        format.set_align('center')
                        format.set_align('vcenter')
                        worksheet = writer.sheets["Vocabulary Test Info #" + str(session_id)]
                        worksheet.set_column('A:A', 13, format)
                        worksheet.set_column('B:B', 16, format)
                        worksheet.set_column('C:C', 18, format)
                        worksheet.set_column('D:D', 16, format)
                        worksheet.set_column('E:E', 16, format)
                        worksheet.set_column('F:I', 16, format)
                        worksheet.set_column('J:K', 20, format)
                        worksheet.set_column('L:L', 28, format)

                if ("results_rc" in results) and ((len(results['results_rc'])) > 0):
                    if rc_total > 0:
                        pd.DataFrame(rc_infos).to_excel(writer, sheet_name="RC Test Info # " + str(session_id),
                                                        index=False)
                        pd.DataFrame([rc_grade], index=["Grade"]).to_excel(writer, startrow=1 + len(rc_infos),
                                                                           sheet_name="RC Test Info # " + str(
                                                                               session_id),
                                                                           header=False)
                        workbook = writer.book
                        format = workbook.add_format()
                        format.set_align('center')
                        format.set_align('vcenter')
                        worksheet = writer.sheets["RC Test Info # " + str(session_id)]
                        worksheet.set_column('A:A', 15, format)
                        worksheet.set_column('B:H', 60, format)
                        worksheet.set_column('I:I', 25, format)
                        worksheet.set_column('J:J', 28, format)

                if ("results_cloze" in results) and ((len(results['results_cloze'])) > 0):
                    if cloze_total > 0:
                        pd.DataFrame(cloze_infos).to_excel(writer, sheet_name="Cloze Test Info # " + str(session_id),
                                                           index=False)
                        pd.DataFrame([cloze_grade], index=["Grade"]).to_excel(writer, startrow=1 + len(cloze_infos),
                                                                              sheet_name="Cloze Test Info # " + str(
                                                                                  session_id),
                                                                              header=False)
                        workbook = writer.book
                        worksheet = writer.sheets["Cloze Test Info # " + str(session_id)]
                        format = workbook.add_format()
                        format.set_align('center')
                        format.set_align('vcenter')
                        worksheet.set_column('A:A', 13, format)
                        worksheet.set_column('B:B', 16, format)
                        worksheet.set_column('C:I', 28, format)

                if ('results_writing' in results) and ((len(results['results_writing'])) > 0):
                    if writing_total > 0:
                        pd.DataFrame(writing_infos).to_excel(writer, sheet_name="Writing Test Info #" + str(session_id),
                                                             index=False)
                        pd.DataFrame([writing_grade], index=["Grade"]).to_excel(writer, startrow=1 + len(writing_infos),
                                                                                sheet_name="Writing Test Info #" + str(
                                                                                    session_id),
                                                                                header=False)
                        workbook = writer.book
                        worksheet = writer.sheets["Writing Test Info #" + str(session_id)]
                        format = workbook.add_format()
                        format.set_align('center')
                        format.set_align('vcenter')
                        worksheet.set_column('A:A', 13, format)
                        worksheet.set_column('B:B', 13, format)
                        worksheet.set_column('C:C', 60, format)
                        worksheet.set_column('D:D', 18, format)
                        worksheet.set_column('E:E', 22, format)
                        worksheet.set_column('F:F', 22, format)
                        worksheet.set_column('G:G', 35, format)
                        # writer.save()

        output = BytesIO()

        with zipfile.ZipFile(output, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
            for session_id in [session.id for session in sessions]:
                t_s = TestSession.query.filter_by(id=session_id).first()
                results = test_schema.dump(t_s)
                if ('results_writing' in results) and ((len(results['results_writing'])) > 0):
                    for i in range(len(writing_results)):
                        txt = BytesIO()
                        docx = BytesIO()
                        if results['results_writing'][i]['answer']['text'] is None:
                            results['results_writing'][i]['answer']['text'] = " "
                        txt.write((results['results_writing'][i]['answer']['text']).encode("utf8"))
                        document = Document()
                        document.add_paragraph(results['results_writing'][i]['answer']['text'])
                        document.save(docx)
                        file_name = "Test #" + str(session_id) + writing_results[i]['name']
                        zip_file.writestr(file_name + ".txt", txt.getvalue())
                        zip_file.writestr(file_name + ".docx", docx.getvalue())
                zip_file.writestr("Test Details user.xlsx", xlsx.getvalue())
            else:
                zip_file.writestr("Test Details user.xlsx", xlsx.getvalue())
            xlsx = BytesIO
            zip_file.writestr("User Test Summary.xlsx",
                              xlsx.getvalue(TestSessionExportProvider().test_session_summary(sessions)))
            zip_file.writestr("Specific Test Summary.xlsx",
                              xlsx.getvalue(TestSessionExportProvider().test_session_summary_for_instructor(sessions)))

        output.seek(0)
        return output

    def test_session_grade_summary(self, results, questions, max_voc_test, max_rc_test, max_cloze_test, max_writing_test):
        voc_questions_test = questions['voc_questions']
        rc_questions_test = questions['rc_questions']
        cloze_questions_test = questions['cloze_questions']
        test_results = results['test']

        test_infos = {
            "User ID": results['user_id'],
            "Test Info": results['name'],
            "User Name": results['user']['name'],
            "First Name": results['user']['first_name'],
            "Last Name": results['user']['last_name']
        }
        if max_voc_test > 0:
            if ('test_vocabulary' in test_results) and (len(test_results['test_vocabulary']) > 0):
                voc_results = test_results["test_vocabulary"]
                for i in range(len(results['results_vocabulary'][0]['answers'])):

                    answer = results['results_vocabulary'][0]['answers'][i]['text']

                    if answer is None:
                        test_result = "Unattempted!" if results['results_vocabulary'][0]['answers'][i][
                                                            'seen'] is True else "Unseen!"
                    elif answer == voc_results[i]['options'][int(voc_results[i]['correct']) - 1]['text']:
                        test_result = 1
                    else:
                        test_result = 0

                    test_infos["voc # " + str(i + 1) + " " + str(voc_questions_test[i])] = test_result

                    if len(results['results_vocabulary'][0]['answers']) < max_voc_test:
                        missing_num_voc = max_voc_test - len(results['results_vocabulary'][0]['answers'])
                        for j in range(missing_num_voc):
                            test_infos["voc #" + str(i + j + 1)] = 0

            else:
                for k in range(max_voc_test):
                    test_infos["voc #" + str(k + 1)]: ""

        if max_rc_test > 0:
            rc_results = test_results["test_rc"]

            if 'test_rc' in test_results and (len(test_results['test_rc']) > 0):

                for i in range(len(rc_results)):
                    rc_questions = test_results["test_rc"][i]['questions']
                    for j in range(len(rc_questions)):
                        answer = results['results_rc'][i]['answers'][j]['text']
                        if answer is None:
                            answer = ""
                        if answer == "":
                            test_result = "Unattempted!" if results['results_rc'][i]['answers'][j][
                                                                'seen'] is True else "Unseen!"
                        elif answer == rc_questions[j]['options'][int(rc_questions[j]['correct']) - 1]['text']:
                            test_result = 1
                        else:
                            test_result = 0
                        test_infos["rc #" + str(j + 1) + " " + str(rc_questions_test[j])] = test_result

                        if len(test_results["test_rc"][i]['questions']) < max_rc_test:
                            missing_num_rc = max_rc_test - len(test_results["test_rc"][i]['questions'])
                            for k in range(missing_num_rc):
                                test_infos["rc #" + str(j + k + 1)]: " "

            else:
                for h in range(max_rc_test):
                    test_infos["rc #" + str(h + 1)]: ""

        if max_cloze_test > 0:
            if 'test_cloze' in test_results and (len(test_results['test_cloze']) > 0):
                cloze_results = test_results["test_cloze"]

                for i in range(len(cloze_results)):
                    cloze_questions = test_results["test_cloze"][i]['questions']
                    for j in range(len(cloze_questions)):
                        answer = results['results_cloze'][i]['answers'][j]['text']
                        if answer is None:
                            test_result = "Unattempted!" if results['results_cloze'][i]['answers'][j][
                                                                'seen'] is True else "Unseen!"
                        elif answer == cloze_questions[j]['options'][int(cloze_questions[j]['correct']) - 1]['text']:
                            test_result = 1
                        else:
                            test_result = 0
                        test_infos["cloze #" + str(j + 1) + " " + str(cloze_questions_test[j])] = test_result

                        if len(cloze_questions) < max_cloze_test:
                            missing_num_cloze = max_cloze_test - len(cloze_questions)
                            for k in range(missing_num_cloze):
                                test_infos["cloze #" + str(j + k + 1)]: " "
            else:
                for h in range(max_cloze_test):
                    test_infos["cloze #" + str(h + 1)]: ""

        if max_writing_test > 0:
            if 'test_writing' in test_results and (len(test_results['test_writing']) > 0):
                writing_results = test_results["test_writing"]
                for i in range(max_writing_test):
                    test_infos["writing # " + str(i + 1)] = 0

                    if len(writing_results) < max_writing_test:
                        missing_num_writing = max_writing_test - len(writing_results)
                        for j in range(missing_num_writing):
                            test_infos["writing # " + str(j + 1)] = " "
            else:
                for h in range(max_cloze_test):
                    test_infos["writing # " + str(h + 1)] = " "

        return test_infos

    def vocabulary_test_grade(self, results):
        if ('results_vocabulary' in results) and ((len(results['results_vocabulary'])) > 0):
            voc_score = testhelper.calculate_correctly_answered_questions(results['results_vocabulary'][0]['answers'])
            voc_total = len(results['results_vocabulary'][0]['answers'])
            voc_grade = testhelper.grade_calculator(voc_score, voc_total)
            if voc_total > 0:
                return voc_grade
        else:
            return None

    def rc_test_grade(self, results):
        if 'results_rc' in results and ((len(results['results_rc'])) > 0):
            rc_score = testhelper.calculate_correctly_answered_questions(results['results_rc'][0]['answers'])
            rc_total = len(results['results_rc'][0]['answers'])
            rc_grade = testhelper.grade_calculator(rc_score, rc_total)
            if rc_total > 0:
                return rc_grade
        else:
            return None

    def cloze_test_grade(self, results):
        if 'results_cloze' in results and ((len(results['results_cloze'])) > 0):
            cloze_score = testhelper.calculate_correctly_answered_questions(results['results_cloze'][0]['answers'])
            cloze_total = len(results['results_cloze'][0]['answers'])
            cloze_grade = testhelper.grade_calculator(cloze_score, cloze_total)
            if cloze_total > 0:
                return cloze_grade
        else:
            return None

    def writing_test_grade(self, results):
        if 'results_writing' in results and ((len(results['results_writing'])) > 0):
            writing_score = testhelper.calculate_correctly_answered_questions_writings(results['results_writing'])
            writing_total = len(results['results_writing'])
            writing_grade = testhelper.grade_calculator(writing_score, writing_total)
            if writing_total > 0:
                return writing_grade
        else:
            return None

    def test_session_summary_for_researcher(self, sessions):
        questions = self.get_questions(sessions)
        records_list = []
        for s in sessions:
            records = {
                "Id": s.id,
                "Name": s.name,
                "Test Id": s.test_id,
                "Test Name": s.test.name,
                "Student Id": s.user_id,
                "Username": s.user.name,
                "First Name": s.user.first_name,
                "Last Name": s.user.last_name,
                "Start Time": s.start_datetime.strftime("%A, %B %d, %Y %I %P"),
                "End Time": s.end_datetime.strftime("%A, %B %d, %Y %I %P"),
                "Created Time": s.created_datetime.strftime("%A, %B %d, %Y %I %P")}
            t_s = TestSession.query.filter_by(id=s.id).first()
            results = test_schema.dump(t_s)
            test_time = testhelper.calculate_time_consumption(results['end_datetime'], results['start_datetime'])
            records["Time Consumption"] = test_time
            records['Vocabulary Grade'] = TestSessionExportProvider().vocabulary_test_grade(results)
            records["Reading Comprehension Grade"] = TestSessionExportProvider().rc_test_grade(results)
            records["Cloze Grade"] = TestSessionExportProvider().cloze_test_grade(results)
            records["Writing Grade"] = "N/A" if TestSessionExportProvider().writing_test_grade(results) is None else TestSessionExportProvider().writing_test_grade(results)
            records_list.append(records)

        output = BytesIO()

        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            pd.DataFrame(records_list).to_excel(writer,
                                                sheet_name="Student Test Summary",
                                                index=False)
            workbook = writer.book
            worksheet = writer.sheets["Student Test Summary"]
            format = workbook.add_format()
            format.set_align('center')
            format.set_align('vcenter')
            worksheet.set_column('A:A', 13, format)
            worksheet.set_column('B:B', 45, format)
            worksheet.set_column('C:C', 13, format)
            worksheet.set_column('D:D', 20, format)
            worksheet.set_column('E:E', 20, format)
            worksheet.set_column('F:F', 20, format)
            worksheet.set_column('G:H', 20, format)
            worksheet.set_column('I:L', 45, format)
            worksheet.set_column('M:P', 36, format)

            max_voc_test = 0
            max_rc_test = 0
            max_cloze_test = 0
            max_writing_test = 0
            for i1 in range(len(sessions)):
                t_s = TestSession.query.filter_by(id=i1 + 1).first()
                result = test_schema.dump(t_s)
                test_results = result['test']
                if 'test_vocabulary' in test_results:
                    voc_results = test_results["test_vocabulary"]
                    if max_voc_test < len(voc_results):
                        max_voc_test = len(voc_results)
                if 'test_rc' in test_results:
                    rc_results = test_results["test_rc"]
                    if max_rc_test < len(rc_results):
                        max_rc_test = len(rc_results)

                if 'test_cloze' in test_results:
                    cloze_results = test_results["test_cloze"]
                    if max_cloze_test < len(cloze_results):
                        max_cloze_test = len(cloze_results)
                if 'test_writing' in test_results:
                    writing_results = test_results["test_writing"]
                    if max_writing_test < len(writing_results):
                        max_writing_test = len(writing_results)

            # for detailed test info
            for i2 in range(len(sessions)):
                t_s = TestSession.query.filter_by(id=i2 + 1).first()
                result = test_schema.dump(t_s)
                test_results = result['test']
                if 'test_rc' in test_results:
                    max_rc_counter = 0
                    for irc in range(len(rc_results)):
                        rc_questions = test_results["test_rc"][irc]['questions']
                        for jrc in range(len(rc_questions)):
                            max_rc_counter = max_rc_counter + 1
                    if max_rc_test < max_rc_counter:
                        max_rc_test = max_rc_counter
                if 'test_cloze' in test_results:
                    max_cloze_counter = 0
                    for icloze in range(len(cloze_results)):
                        cloze_questions = test_results["test_cloze"][icloze]['questions']
                        for jcloze in range(len(cloze_questions)):
                            max_cloze_counter = max_cloze_counter + 1
                    if max_cloze_test < max_cloze_counter:
                        max_cloze_test = max_cloze_counter
            test_info_with_zero_one_result = []
            for i in range(len(sessions)):
                t_s = TestSession.query.filter_by(id=i + 1).first()
                result = test_schema.dump(t_s)

                test_info_with_zero_one_result.append(TestSessionExportProvider().test_session_grade_summary(result, questions, max_voc_test,
                                                                                        max_rc_test, max_cloze_test,
                                                                                        max_writing_test))
            pd.DataFrame(test_info_with_zero_one_result).to_excel(writer,
                                             sheet_name="Test Results",
                                             index=False)
            workbook = writer.book
            format = workbook.add_format()
            format.set_align('center')
            format.set_align('vcenter')
            worksheet = writer.sheets["Test Results"]
            worksheet.set_column('A:A', 13, format)
            worksheet.set_column('B:B', 45, format)
            worksheet.set_column('C:E', 16, format)
            worksheet.set_column('F:AD', 20, format)

            test_info_with_user_answer = []
            for i in range(len(sessions)):
                t_s = TestSession.query.filter_by(id=i + 1).first()
                result = test_schema.dump(t_s)

                test_info_with_user_answer.append(
                    TestSessionExportProvider().test_session_user_answer_summary(result, questions, max_voc_test,
                                                                                 max_rc_test,
                                                                                 max_cloze_test,
                                                                                 max_writing_test))
            pd.DataFrame(test_info_with_user_answer).to_excel(writer,
                                             sheet_name="Student Answer",
                                             index=False)
            workbook = writer.book
            format = workbook.add_format()
            format.set_align('center')
            format.set_align('vcenter')
            worksheet = writer.sheets["Student Answer"]
            worksheet.set_column('A:A', 13, format)
            worksheet.set_column('B:B', 45, format)
            worksheet.set_column('C:E', 16, format)
            worksheet.set_column('F:AD', 20, format)
            test_info_with_correct_answer = []
            for i in range(len(sessions)):
                t_s = TestSession.query.filter_by(id=i + 1).first()
                result = test_schema.dump(t_s)

                test_info_with_correct_answer.append(
                    TestSessionExportProvider().test_session_correct_answer_summary(result, questions, max_voc_test,
                                                                                 max_rc_test,
                                                                                 max_cloze_test,
                                                                                 max_writing_test))
            pd.DataFrame(test_info_with_correct_answer).to_excel(writer,
                                                              sheet_name="Correct Answer",
                                                              index=False)
            workbook = writer.book
            format = workbook.add_format()
            format.set_align('center')
            format.set_align('vcenter')
            worksheet = writer.sheets["Correct Answer"]
            worksheet.set_column('A:A', 13, format)
            worksheet.set_column('B:B', 45, format)
            worksheet.set_column('C:E', 16, format)
            worksheet.set_column('F:AD', 20, format)
            writer.save()

        output.seek(0)
        return output

    def test_session_user_answer_summary(self, results, questions, max_voc_test, max_rc_test, max_cloze_test, max_writing_test):
        voc_questions_test = questions['voc_questions']
        rc_questions_test = questions['rc_questions']
        cloze_questions_test = questions['cloze_questions']
        test_results = results['test']


        test_infos = {
            "User ID": results['user_id'],
            "Test Info": results['name'],
            "User Name": results['user']['name'],
            "First Name": results['user']['first_name'],
            "Last Name": results['user']['last_name']
        }
        if max_voc_test > 0:
            if ('test_vocabulary' in test_results) and (len(test_results['test_vocabulary']) > 0):
                voc_results = test_results["test_vocabulary"]
                for i in range(len(results['results_vocabulary'][0]['answers'])):

                    answer = results['results_vocabulary'][0]['answers'][i]['text']

                    if answer is None:
                        test_result = "Unattempted!" if results['results_vocabulary'][0]['answers'][i][
                                                            'seen'] is True else "Unseen!"
                    else:
                        test_result = answer

                    test_infos["voc # " + str(i + 1) + " " + str(voc_questions_test[i])] = test_result

                    if len(results['results_vocabulary'][0]['answers']) < max_voc_test:
                        missing_num_voc = max_voc_test - len(results['results_vocabulary'][0]['answers'])
                        for j in range(missing_num_voc):
                            test_infos["voc #" + str(i + j + 1)] = ""

            else:
                for k in range(max_voc_test):
                    test_infos["voc #" + str(k + 1)]: ""

        if max_rc_test > 0:
            rc_results = test_results["test_rc"]

            if 'test_rc' in test_results and (len(test_results['test_rc']) > 0):

                for i in range(len(rc_results)):
                    rc_questions = test_results["test_rc"][i]['questions']
                    for j in range(len(rc_questions)):
                        answer = results['results_rc'][i]['answers'][j]['text']
                        if answer is None:
                            test_result = "Unattempted!" if results['results_rc'][i]['answers'][j][
                                                                'seen'] is True else "Unseen!"
                        else:
                            test_result = answer
                        test_infos["rc #" + str(j + 1) + " " + str(rc_questions_test[j])] = test_result

                        if len(test_results["test_rc"][i]['questions']) < max_rc_test:
                            missing_num_rc = max_rc_test - len(test_results["test_rc"][i]['questions'])
                            for k in range(missing_num_rc):
                                test_infos["rc #" + str(j + k + 1)]: " "

            else:
                for h in range(max_rc_test):
                    test_infos["rc #" + str(h + 1)]: ""

        if max_cloze_test > 0:
            if 'test_cloze' in test_results and (len(test_results['test_cloze']) > 0):
                cloze_results = test_results["test_cloze"]

                for i in range(len(cloze_results)):
                    cloze_questions = test_results["test_cloze"][i]['questions']
                    for j in range(len(cloze_questions)):
                        answer = results['results_cloze'][i]['answers'][j]['text']
                        if answer is None:
                            test_result = "Unattempted!" if results['results_cloze'][i]['answers'][j][
                                                                'seen'] is True else "Unseen!"
                        else:
                            test_result = answer
                        test_infos["cloze #" + str(j + 1) + " " + str(cloze_questions_test[j])] = test_result

                        if len(cloze_questions) < max_cloze_test:
                            missing_num_cloze = max_cloze_test - len(cloze_questions)
                            for k in range(missing_num_cloze):
                                test_infos["cloze #" + str(j + k + 1)]: " "
            else:
                for h in range(max_cloze_test):
                    test_infos["cloze #" + str(h + 1)]: ""

        if max_writing_test > 0:
            if 'test_writing' in test_results and (len(test_results['test_writing']) > 0):
                writing_results = test_results["test_writing"]
                for i in range(max_writing_test):
                    test_infos["writing # " + str(i + 1)] = " "

                    if len(writing_results) < max_writing_test:
                        missing_num_writing = max_writing_test - len(writing_results)
                        for j in range(missing_num_writing):
                            test_infos["writing # " + str(j + 1)] = " "
            else:
                for h in range(max_cloze_test):
                    test_infos["writing # " + str(h + 1)] = " "

        return test_infos

    def test_session_summary_for_instructor(self, sessions):
        questions = self.get_questions(sessions)
        records_list = []
        for s in sessions:
            records = {
                "Id": s.id,
                "Name": s.name,
                "Test Id": s.test_id,
                "Test Name": s.test.name,
                "Student Id": s.user_id,
                "Username": s.user.name,
                "First Name": s.user.first_name,
                "Last Name": s.user.last_name,
                "Start Time": s.start_datetime.strftime("%A, %B %d, %Y %I %P"),
                "End Time": s.end_datetime.strftime("%A, %B %d, %Y %I %P"),
                "Created Time": s.created_datetime.strftime("%A, %B %d, %Y %I %P")}
            t_s = TestSession.query.filter_by(id=s.id).first()
            results = test_schema.dump(t_s)

            test_time = testhelper.calculate_time_consumption(results['end_datetime'], results['start_datetime'])
            records["Time Consumption"] = test_time
            records['Vocabulary Grade'] = TestSessionExportProvider().vocabulary_test_grade(results)
            records["Reading Comprehension Grade"] = TestSessionExportProvider().rc_test_grade(results)
            records["Cloze Grade"] = TestSessionExportProvider().cloze_test_grade(results)
            records["Writing Grade"] = "N/A" if TestSessionExportProvider().writing_test_grade(results) is None else TestSessionExportProvider().writing_test_grade(results)
            records_list.append(records)

        output = BytesIO()

        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            pd.DataFrame(records_list).to_excel(writer,
                                                sheet_name="Student Test Summary",
                                                index=False)
            workbook = writer.book
            worksheet = writer.sheets["Student Test Summary"]
            format = workbook.add_format()
            format.set_align('center')
            format.set_align('vcenter')
            worksheet.set_column('A:A', 13, format)
            worksheet.set_column('B:B', 45, format)
            worksheet.set_column('C:C', 13, format)
            worksheet.set_column('D:D', 20, format)
            worksheet.set_column('E:E', 20, format)
            worksheet.set_column('F:F', 20, format)
            worksheet.set_column('G:H', 20, format)
            worksheet.set_column('I:L', 45, format)
            worksheet.set_column('M:P', 36, format)

            max_voc_test = 0
            max_rc_test = 0
            max_cloze_test = 0
            max_writing_test = 0
            for i1 in range(len(sessions)):
                t_s = TestSession.query.filter_by(id=i1 + 1).first()
                result = test_schema.dump(t_s)
                test_results = result['test']
                if 'test_vocabulary' in test_results:
                    voc_results = test_results["test_vocabulary"]
                    if max_voc_test < len(voc_results):
                        max_voc_test = len(voc_results)
                if 'test_rc' in test_results:
                    rc_results = test_results["test_rc"]
                    if max_rc_test < len(rc_results):
                        max_rc_test = len(rc_results)

                if 'test_cloze' in test_results:
                    cloze_results = test_results["test_cloze"]
                    if max_cloze_test < len(cloze_results):
                        max_cloze_test = len(cloze_results)
                if 'test_writing' in test_results:
                    writing_results = test_results["test_writing"]
                    if max_writing_test < len(writing_results):
                        max_writing_test = len(writing_results)

            # for detailed test info
            for i2 in range(len(sessions)):
                t_s = TestSession.query.filter_by(id=i2 + 1).first()
                result = test_schema.dump(t_s)
                test_results = result['test']
                if 'test_rc' in test_results:
                    max_rc_counter = 0
                    for irc in range(len(rc_results)):
                        rc_questions = test_results["test_rc"][irc]['questions']
                        for jrc in range(len(rc_questions)):
                            max_rc_counter = max_rc_counter + 1
                    if max_rc_test < max_rc_counter:
                        max_rc_test = max_rc_counter
                if 'test_cloze' in test_results:
                    max_cloze_counter = 0
                    for icloze in range(len(cloze_results)):
                        cloze_questions = test_results["test_cloze"][icloze]['questions']
                        for jcloze in range(len(cloze_questions)):
                            max_cloze_counter = max_cloze_counter + 1
                    if max_cloze_test < max_cloze_counter:
                        max_cloze_test = max_cloze_counter
            test_info = []
            for i in range(len(sessions)):
                t_s = TestSession.query.filter_by(id=i + 1).first()
                result = test_schema.dump(t_s)

                test_info.append(TestSessionExportProvider().test_session_user_answer_summary(result, questions, max_voc_test,
                                                                                              max_rc_test,
                                                                                              max_cloze_test,
                                                                                              max_writing_test))
            pd.DataFrame(test_info).to_excel(writer,
                                             sheet_name="Student Answer",
                                             index=False)
            workbook = writer.book
            format = workbook.add_format()
            format.set_align('center')
            format.set_align('vcenter')
            worksheet = writer.sheets["Student Answer"]
            worksheet.set_column('A:A', 13, format)
            worksheet.set_column('B:B', 45, format)
            worksheet.set_column('C:E', 16, format)
            worksheet.set_column('F:AD', 20, format)
            writer.save()

        output.seek(0)
        return output

    def write_results_into_file_researcher_instructor(self, sessions, name):
        xlsx = BytesIO()
        with pd.ExcelWriter(xlsx, engine='xlsxwriter') as writer:
            for session_id in [session.id for session in sessions]:
                t_s = TestSession.query.filter_by(id=session_id).first()
                results = test_schema.dump(t_s)
                test_time = testhelper.calculate_time_consumption(results['end_datetime'], results['start_datetime'])

                basic_info = {
                    "User ID": results['user_id'],
                    "Test Info": results['name'],
                    "User Name": results['user']['name'],
                    "First Name": results['user']['first_name'],
                    "Last Name": results['user']['last_name'],
                    "Start Time": results['start_datetime'],
                    "End Time": results['end_datetime'],
                    "Time Consumption": test_time
                }

                if ('results_vocabulary' in results) and ((len(results['results_vocabulary'])) > 0):
                    voc_score = testhelper.calculate_correctly_answered_questions(
                        results['results_vocabulary'][0]['answers'])
                    voc_total = len(results['results_vocabulary'][0]['answers'])
                    voc_grade = testhelper.grade_calculator(voc_score, voc_total)
                    if voc_total > 0:
                        basic_info["Vocabulary Grade"] = voc_grade

                if 'results_rc' in results and ((len(results['results_rc'])) > 0):
                    rc_score = testhelper.calculate_correctly_answered_questions(results['results_rc'][0]['answers'])
                    rc_total = len(results['results_rc'][0]['answers'])
                    rc_grade = testhelper.grade_calculator(rc_score, rc_total)
                    if rc_total > 0:
                        basic_info["Reading Comprehension Grade"] = rc_grade

                if 'results_cloze' in results and ((len(results['results_cloze'])) > 0):
                    cloze_score = testhelper.calculate_correctly_answered_questions(
                        results['results_cloze'][0]['answers'])
                    cloze_total = len(results['results_cloze'][0]['answers'])
                    cloze_grade = testhelper.grade_calculator(cloze_score, cloze_total)
                    if cloze_total > 0:
                        basic_info["Cloze Grade"] = cloze_grade

                if 'results_writing' in results and ((len(results['results_writing'])) > 0):
                    writing_score = testhelper.calculate_correctly_answered_questions_writings(
                        results['results_writing'])
                    writing_total = len(results['results_writing'])
                    writing_grade = testhelper.grade_calculator(writing_score, writing_total)
                    if writing_total > 0:
                        basic_info["Writing Grade"] = writing_grade
                else:
                    basic_info["Writing Grade"] = None

                test_results = results['test']

                if ('test_vocabulary' in test_results) and (len(test_results['test_vocabulary']) > 0):
                    voc_results = test_results["test_vocabulary"]
                    voc_infos = []
                    for i in range(len(results['results_vocabulary'][0]['answers'])):
                        answer = results['results_vocabulary'][0]['answers'][i]['text']
                        if answer is None:
                            test_result = "Empty Answer"
                        elif answer == voc_results[i]['options'][int(voc_results[i]['correct']) - 1]['text']:
                            test_result = "Correct"
                        else:
                            test_result = "Incorrect"
                        test_voc_time = testhelper.calculate_time_consumption(
                            results['results_vocabulary'][0]['answers'][i]['end_time'],
                            results['results_vocabulary'][0]['answers'][i]['start_time'])
                        voc_infos.append({
                            "Vocabulary": str(i + 1),
                            "Test Type": voc_results[i]['type'],
                            "Word": voc_results[i]['word'],
                            "Option 1": voc_results[i]['options'][0]['text'],
                            "Option 2": voc_results[i]['options'][1]['text'],
                            "Option 3": voc_results[i]['options'][2]['text'],
                            "Option 4": voc_results[i]['options'][3]['text'],
                            "Student Answer": answer,
                            "Correct Answer": voc_results[i]['options'][int(voc_results[i]['correct']) - 1]['text'],
                            "Difficulty": voc_results[i]['difficulty'],
                            "Test Result": test_result,
                            "Time Consumption": test_voc_time
                        })

                if 'test_rc' in test_results and (len(test_results['test_rc']) > 0):
                    rc_results = test_results["test_rc"]
                    rc_infos = []
                    for i in range(len(rc_results)):
                        rc_questions = test_results["test_rc"][i]['questions']
                        for j in range(len(rc_questions)):
                            answer = results['results_rc'][i]['answers'][j]['text']
                            if answer is None:
                                test_result = "Empty Answer"
                            elif answer == rc_questions[j]['options'][int(rc_questions[j]['correct']) - 1]['text']:
                                test_result = "Correct"
                            else:
                                test_result = "Incorrect"
                            test_rc_time = testhelper.calculate_time_consumption(
                                results['results_rc'][i]['answers'][j]['end_time'],
                                results['results_rc'][i]['answers'][j]['start_time'])
                            rc_infos.append({
                                "RC": str(j + 1),
                                "Question": rc_questions[j]['text'],
                                "Option 1": rc_questions[j]['options'][0]['text'],
                                "Option 2": rc_questions[j]['options'][1]['text'],
                                "Option 3": rc_questions[j]['options'][2]['text'],
                                "Option 4": rc_questions[j]['options'][3]['text'],
                                "User Answer": results['results_rc'][i]['answers'][j]['text'],
                                "Correct Answer": rc_questions[j]['options'][int(rc_questions[j]['correct']) - 1][
                                    'text'],
                                "Test Result": test_result,
                                "Time Consumption": test_rc_time
                            })

                if 'test_cloze' in test_results and (len(test_results['test_cloze']) > 0):
                    cloze_results = test_results["test_cloze"]

                    cloze_infos = []
                    for i in range(len(cloze_results)):
                        cloze_questions = test_results["test_cloze"][i]['questions']
                        for j in range(len(cloze_questions)):
                            answer = results['results_cloze'][i]['answers'][j]['text']
                            if answer is None:
                                test_result = "Empty Answer"
                            elif answer == cloze_questions[j]['options'][int(cloze_questions[j]['correct']) - 1][
                                'text']:
                                test_result = "Correct"
                            else:
                                test_result = "Incorrect"
                            test_cloze_time = testhelper.calculate_time_consumption(
                                results['results_cloze'][i]['answers'][j]['end_time'],
                                results['results_cloze'][i]['answers'][j]['start_time'])
                            cloze_infos.append({
                                "Cloze": str(j + 1),
                                "Question": cloze_questions[j]['text'],
                                "Option 1": cloze_questions[j]['options'][0]['text'],
                                "Option 2": cloze_questions[j]['options'][1]['text'],
                                "Option 3": cloze_questions[j]['options'][2]['text'],
                                "Option 4": cloze_questions[j]['options'][3]['text'],
                                "User Answer": results['results_cloze'][i]['answers'][j]['text'],
                                "Correct Answer": cloze_questions[j]['options'][int(cloze_questions[j]['correct']) - 1][
                                    'text'],
                                "Test Result": test_result,
                                "Time Consumption": test_cloze_time
                            })

                if 'test_writing' in test_results and (len(test_results['test_writing']) > 0):
                    writing_results = test_results["test_writing"]
                    writing_infos = []
                    for i in range(len(writing_results)):
                        test_writing_time = testhelper.calculate_time_consumption(
                            results['results_writing'][i]['answer']['end_time'],
                            results['results_writing'][i]['answer']['start_time'])
                        writing_infos.append({
                            "Writing": str(i + 1),
                            "Test name": writing_results[i]['name'],
                            "Question": writing_results[i]['question'],
                            "Word Limit": writing_results[i]['word_limit'],
                            "Time Limit": writing_results[i]['time_limit'],
                            "Essay": results['results_writing'][i]['answer']['text'],
                            "Time Consumption": test_writing_time
                        })

                if ('results_vocabulary' in results) and ((len(results['results_vocabulary'])) > 0):
                    if voc_total > 0:
                        pd.DataFrame(voc_infos).to_excel(writer, sheet_name="Vocabulary Test Info #" + str(session_id),
                                                         index=False)
                        pd.DataFrame([voc_grade], index=["Grade"]).to_excel(writer, startrow=1 + len(voc_infos),
                                                                            sheet_name="Vocabulary Test Info #" + str(
                                                                                session_id),
                                                                            header=False)
                        workbook = writer.book
                        format = workbook.add_format()
                        format.set_align('center')
                        format.set_align('vcenter')
                        worksheet = writer.sheets["Vocabulary Test Info #" + str(session_id)]
                        worksheet.set_column('A:A', 13, format)
                        worksheet.set_column('B:B', 16, format)
                        worksheet.set_column('C:C', 18, format)
                        worksheet.set_column('D:D', 16, format)
                        worksheet.set_column('E:E', 16, format)
                        worksheet.set_column('F:I', 16, format)
                        worksheet.set_column('J:K', 20, format)
                        worksheet.set_column('L:L', 28, format)

                if ("results_rc" in results) and ((len(results['results_rc'])) > 0):
                    if rc_total > 0:
                        pd.DataFrame(rc_infos).to_excel(writer, sheet_name="RC Test Info # " + str(session_id),
                                                        index=False)
                        pd.DataFrame([rc_grade], index=["Grade"]).to_excel(writer, startrow=1 + len(rc_infos),
                                                                           sheet_name="RC Test Info # " + str(
                                                                               session_id),
                                                                           header=False)
                        workbook = writer.book
                        format = workbook.add_format()
                        format.set_align('center')
                        format.set_align('vcenter')
                        worksheet = writer.sheets["RC Test Info # " + str(session_id)]
                        worksheet.set_column('A:A', 15, format)
                        worksheet.set_column('B:H', 60, format)
                        worksheet.set_column('I:I', 25, format)
                        worksheet.set_column('J:J', 28, format)

                if ("results_cloze" in results) and ((len(results['results_cloze'])) > 0):
                    if cloze_total > 0:
                        pd.DataFrame(cloze_infos).to_excel(writer, sheet_name="Cloze Test Info # " + str(session_id),
                                                           index=False)
                        pd.DataFrame([cloze_grade], index=["Grade"]).to_excel(writer, startrow=1 + len(cloze_infos),
                                                                              sheet_name="Cloze Test Info # " + str(
                                                                                  session_id),
                                                                              header=False)
                        workbook = writer.book
                        worksheet = writer.sheets["Cloze Test Info # " + str(session_id)]
                        format = workbook.add_format()
                        format.set_align('center')
                        format.set_align('vcenter')
                        worksheet.set_column('A:A', 13, format)
                        worksheet.set_column('B:B', 16, format)
                        worksheet.set_column('C:I', 28, format)

                if ('results_writing' in results) and ((len(results['results_writing'])) > 0):
                    if writing_total > 0:
                        pd.DataFrame(writing_infos).to_excel(writer, sheet_name="Writing Test Info #" + str(session_id),
                                                             index=False)
                        pd.DataFrame([writing_grade], index=["Grade"]).to_excel(writer, startrow=1 + len(writing_infos),
                                                                                sheet_name="Writing Test Info #" + str(
                                                                                    session_id),
                                                                                header=False)
                        workbook = writer.book
                        worksheet = writer.sheets["Writing Test Info #" + str(session_id)]
                        format = workbook.add_format()
                        format.set_align('center')
                        format.set_align('vcenter')
                        worksheet.set_column('A:A', 13, format)
                        worksheet.set_column('B:B', 13, format)
                        worksheet.set_column('C:C', 60, format)
                        worksheet.set_column('D:D', 18, format)
                        worksheet.set_column('E:E', 22, format)
                        worksheet.set_column('F:F', 22, format)
                        worksheet.set_column('G:G', 35, format)
                        # writer.save()

        output = BytesIO()

        with zipfile.ZipFile(output, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
            for session_id in [session.id for session in sessions]:
                t_s = TestSession.query.filter_by(id=session_id).first()
                results = test_schema.dump(t_s)
                if ('results_writing' in results) and ((len(results['results_writing'])) > 0):
                    for i in range(len(writing_results)):
                        txt = BytesIO()
                        docx = BytesIO()
                        if results['results_writing'][i]['answer']['text'] is None:
                            results['results_writing'][i]['answer']['text'] = " "
                        txt.write((results['results_writing'][i]['answer']['text']).encode("utf8"))
                        document = Document()
                        document.add_paragraph(results['results_writing'][i]['answer']['text'])
                        document.save(docx)
                        file_name = "Test #" + str(session_id) + writing_results[i]['name']
                        zip_file.writestr(file_name + ".txt", txt.getvalue())
                        zip_file.writestr(file_name + ".docx", docx.getvalue())
                #zip_file.writestr("Test Details test.xlsx", xlsx.getvalue())


            else:
                pass
                #zip_file.writestr("Test Details test.xlsx", xlsx.getvalue())
            xlsx = BytesIO
            zip_file.writestr("Test results for researcher.xlsx",
                              xlsx.getvalue(TestSessionExportProvider().test_session_summary_for_researcher(sessions)))
            zip_file.writestr("Test results for instructor.xlsx",
                              xlsx.getvalue(TestSessionExportProvider().test_session_summary_for_instructor(sessions)))

        output.seek(0)
        return output

    def get_questions(self, sessions):
        session_id = sessions[0].id
        t_s = TestSession.query.filter_by(id=session_id).first()
        results = test_schema.dump(t_s)
        voc_questions = []
        voc_results = results['test']['test_vocabulary']
        for voc_question in voc_results:
            voc_questions.append(voc_question['word'])

        rc_questions = []

        rc_results = results['test']['test_rc'][0]['questions']
        for rc_question in rc_results:
            rc_questions.append(rc_question['text'])

        cloze_questions = []
        cloze_results = results['test']['test_cloze'][0]['questions']
        for cloze_question in cloze_results:
            cloze_questions.append(cloze_question['text'])

        return {'voc_questions': voc_questions, 'rc_questions': rc_questions, 'cloze_questions': cloze_questions}

    def test_session_correct_answer_summary(self, results, questions, max_voc_test, max_rc_test, max_cloze_test, max_writing_test):
        voc_questions_test = questions['voc_questions']
        rc_questions_test = questions['rc_questions']
        cloze_questions_test = questions['cloze_questions']
        test_results = results['test']
        test_infos = {
            "User ID": results['user_id'],
            "Test Info": results['name'],
            "User Name": results['user']['name'],
            "First Name": results['user']['first_name'],
            "Last Name": results['user']['last_name']
        }

        if max_voc_test > 0:
            if ('test_vocabulary' in test_results) and (len(test_results['test_vocabulary']) > 0):
                voc_results = test_results["test_vocabulary"]
                for i in range(len(results['results_vocabulary'][0]['answers'])):

                    answer = results['test']['test_vocabulary'][i]['word']
                    test_infos["voc # " + str(i + 1) + " " + str(voc_questions_test[i])] = answer

                    if len(results['results_vocabulary'][0]['answers']) < max_voc_test:
                        missing_num_voc = max_voc_test - len(results['results_vocabulary'][0]['answers'])
                        for j in range(missing_num_voc):
                            test_infos["voc #" + str(i + j + 1)] = ""

            else:
                for k in range(max_voc_test):
                    test_infos["voc #" + str(k + 1)]: ""
        if max_rc_test > 0:
            rc_results = test_results["test_rc"]

            if 'test_rc' in test_results and (len(test_results['test_rc']) > 0):

                for i in range(len(rc_results)):
                    rc_questions = test_results["test_rc"][i]['questions']
                    for j in range(len(rc_questions)):
                        correct_option = int(results['test']['test_rc'][i]['questions'][j]['correct'])
                        answer = results['test']['test_rc'][i]['questions'][j]['options'][correct_option-1]['text']

                        test_infos["rc #" + str(j + 1) + " " + str(rc_questions_test[j])] = answer

                        if len(test_results["test_rc"][i]['questions']) < max_rc_test:
                            missing_num_rc = max_rc_test - len(test_results["test_rc"][i]['questions'])
                            for k in range(missing_num_rc):
                                test_infos["rc #" + str(j + k + 1)]: " "

            else:
                for h in range(max_rc_test):
                    test_infos["rc #" + str(h + 1)]: ""

        if max_cloze_test > 0:
            if 'test_cloze' in test_results and (len(test_results['test_cloze']) > 0):
                cloze_results = test_results["test_cloze"]

                for i in range(len(cloze_results)):
                    cloze_questions = test_results["test_cloze"][i]['questions']
                    for j in range(len(cloze_questions)):
                        answer = results['test']['test_cloze'][i]['questions'][j]['text']

                        test_infos["cloze #" + str(j + 1) + " " + str(cloze_questions_test[j])] = answer

                        if len(cloze_questions) < max_cloze_test:
                            missing_num_cloze = max_cloze_test - len(cloze_questions)
                            for k in range(missing_num_cloze):
                                test_infos["cloze #" + str(j + k + 1)]: " "
            else:
                for h in range(max_cloze_test):
                    test_infos["cloze #" + str(h + 1)]: ""

        if max_writing_test > 0:
            if 'test_writing' in test_results and (len(test_results['test_writing']) > 0):
                writing_results = test_results["test_writing"]
                for i in range(max_writing_test):
                    test_infos["writing # " + str(i + 1)] = " "

                    if len(writing_results) < max_writing_test:
                        missing_num_writing = max_writing_test - len(writing_results)
                        for j in range(missing_num_writing):
                            test_infos["writing # " + str(j + 1)] = " "
            else:
                for h in range(max_cloze_test):
                    test_infos["writing # " + str(h + 1)] = " "

        return test_infos

